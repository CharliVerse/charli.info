from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
BLACK = RGBColor(0, 0, 0)


def set_font(run, name="Arial", size=11, bold=None, italic=None, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc, preset="standard"):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10 if preset == "standard" else 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 16 if preset == "standard" else 18, 8 if preset == "standard" else 10),
        "Heading 2": (13, BLUE, 12 if preset == "standard" else 14, 6 if preset == "standard" else 7),
        "Heading 3": (12, NAVY, 8 if preset == "standard" else 10, 4 if preset == "standard" else 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5 if preset == "standard" else 0.375)
        style.paragraph_format.first_line_indent = Inches(-0.25 if preset == "standard" else -0.188)
        style.paragraph_format.space_after = Pt(8 if preset == "standard" else 4)
        style.paragraph_format.line_spacing = 1.167 if preset == "standard" else 1.25

    header = section.header.paragraphs[0]
    header.text = "ZERO INFORMATION CASUALTY"
    set_font(header.runs[0], size=9, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_title_block(doc, title, subtitle, version_line):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_font(run, size=23, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(subtitle)
    set_font(run, size=13, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(version_line)
    set_font(run, size=10, bold=True, color=BLUE)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def set_core_properties(doc, title, subject):
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Charlotte Joanne Tyrer"
    props.keywords = "Zero Information Casualty, ZIC, accessibility, AI agents, open web"
    props.comments = "Licensed under CC BY-SA 4.0."


def build_accessibility_statement():
    doc = Document()
    configure_document(doc, "standard")
    set_core_properties(doc, "Accessibility Statement - Charli.info", "Human and agent access commitments for Charli.info")
    add_title_block(
        doc,
        "Accessibility Statement",
        "Charli.info: access for humans, assistive technologies, and AI agents",
        "Revised for the ZIC-Open v1.1 posture | 11 July 2026",
    )

    p = doc.add_paragraph()
    r = p.add_run("Our commitment. ")
    set_font(r, bold=True, color=NAVY)
    p.add_run(
        "Charli.info follows the Zero Information Casualty baseline. Information must remain available without relying on sight, precision pointing, motion tolerance, default visual settings, or a particular kind of reader."
    )

    doc.add_heading("What this site commits to", level=1)
    for item in (
        "Core content is delivered as real text in semantic HTML and remains available without JavaScript.",
        "Every page provides a skip link, a main landmark, clear heading structure, descriptive links, and keyboard-visible focus.",
        "The site respects operating-system preferences for reduced motion, colour scheme, text size, and increased contrast.",
        "Links and states remain identifiable without relying on colour alone.",
        "Page titles, descriptions, canonical URLs, structured data, licence metadata, and the sitemap are maintained as part of the access surface.",
        "Any informative image must include alternative text that conveys its purpose and relevant meaning. Purely decorative images use empty alternative text or are omitted.",
    ):
        add_bullet(doc, item)

    doc.add_heading("Access for AI agents", level=1)
    doc.add_paragraph(
        "Charli.info adopts the ZIC-Open posture. Public information is affirmatively open to machine readers, not merely left unblocked. The site welcomes training crawlers, search and retrieval crawlers, and user-directed fetchers. This includes agents acting on behalf of people who ask them to find, read, compare, or summarise public material."
    )
    doc.add_paragraph(
        "The crawler welcome is stated in robots.txt. The curated site map for agents is published in llms.txt, with a fuller text rendering in llms-full.txt. The public GitHub repository is a second front door: it exposes the source and build configuration as well as the rendered site."
    )
    doc.add_paragraph(
        "The content is licensed under Creative Commons Attribution-ShareAlike 4.0 International. The ZIC-Open training position follows that licence: sharing and adaptation are permitted with attribution and share-alike obligations, so the public agent policy admits training, search, and user-directed crawler classes."
    )
    for label, url in (
        ("Agent map", "https://charli.info/llms.txt"),
        ("Crawler policy", "https://charli.info/robots.txt"),
        ("Public source repository", "https://github.com/CharliVerse/charli.info"),
        ("Content licence", "https://creativecommons.org/licenses/by-sa/4.0/"),
    ):
        p = doc.add_paragraph()
        r = p.add_run(f"{label}: ")
        set_font(r, bold=True, color=NAVY)
        p.add_run(url)

    doc.add_heading("Current scope", level=1)
    doc.add_paragraph(
        "The public site is intentionally small and uses shared navigation, a shared semantic layout, and a shared stylesheet. Its simplicity is part of the accessibility strategy: the source order is the reading order, and content is not hidden behind interaction or presentation."
    )

    doc.add_heading("Verification", level=1)
    doc.add_paragraph(
        "The site is checked in source and after its static build. Automated checks cover Astro diagnostics, heading order, landmarks, metadata, crawler files, licence declarations, image alternative text, the sitemap, and the absence of anti-agent directives."
    )
    doc.add_paragraph(
        "Verification also has to happen from outside the repository. A plain client must be able to retrieve the public pages without JavaScript, cookies, a browser exception, or a bot challenge. At least one live AI assistant should be asked to fetch and summarise a page by URL; its account must reflect the actual page. This catches failures in DNS, TLS, hosting, or crawler controls that an HTML-only audit cannot see."
    )

    doc.add_heading("Known limits and ongoing work", level=1)
    doc.add_paragraph(
        "Automated checks cannot prove that every experience is accessible. The agent surface is being moved toward build-time generation so that the HTML, Markdown, llms files, and sitemap cannot drift apart. Serial content will provide an RSS or Atom feed when it is published on this domain."
    )

    doc.add_heading("Feedback", level=1)
    doc.add_paragraph(
        "If any part of Charli.info becomes difficult to navigate, retrieve, or understand with assistive technology or an AI agent, that is a defect rather than an edge case. Please report the page, the technology or agent used, what you expected, and what happened."
    )
    doc.add_paragraph("Charlotte Joanne Tyrer | Charli.info | Content licensed CC BY-SA 4.0")

    path = OUT / "Charli.info_Accessibility_Statement_ZIC-Open_v1.1.docx"
    doc.save(path)
    return path


def build_changelog():
    doc = Document()
    configure_document(doc, "compact")
    set_core_properties(doc, "ZIC v1.2 Proposed Changelog", "Proposed corrections and clarifications following the Charli.info field audit")
    add_title_block(
        doc,
        "Zero Information Casualty v1.2",
        "Proposed changelog and specification corrections",
        "Draft for review | 11 July 2026 | Builds on ZIC v1.1",
    )

    p = doc.add_paragraph()
    r = p.add_run("Status. ")
    set_font(r, bold=True, color=NAVY)
    p.add_run(
        "This is a proposed changelog, not a declaration that v1.2 has been adopted. It records corrections and clarifications revealed by the first end-to-end audit of Charli.info against ZIC-Open v1.1."
    )

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Version 1.2 should preserve the philosophy and compliance threshold of v1.1 while correcting two internal contradictions, making the feed rule explicitly conditional, and strengthening outside-in verification of the delivery stack."
    )

    doc.add_heading("Proposed normative changes", level=1)

    doc.add_heading("1. Complete the reference crawler configuration", level=2)
    doc.add_paragraph(
        "Part 3.6 names Meta-ExternalAgent as a training crawler, but the reference robots.txt block omits it. Add the missing user-agent entry so the prose and reference configuration agree."
    )
    add_bullet(doc, "Add: User-agent: Meta-ExternalAgent, followed by Allow: /.")
    add_bullet(doc, "Retain the explanatory comments that make the welcome explicit and identify the three crawler classes.")

    doc.add_heading("2. Clarify focus indicators for programmatic landmark targets", level=2)
    doc.add_paragraph(
        "Part 2.1 currently says that removing or weakening a focus indicator is never acceptable. Add a narrow exception for a landmark that receives focus only as the programmatic destination of a skip link."
    )
    add_bullet(doc, "Visible focus remains mandatory for every user-operable control and link.")
    add_bullet(doc, "A programmatically focused main landmark may suppress its own outline when the skip link remains visibly focused before activation and no interactive element loses a focus indicator.")
    add_bullet(doc, "The exception must not be used for controls, links, form fields, or elements reachable through ordinary sequential keyboard navigation.")

    doc.add_heading("3. Make the syndication-feed requirement explicitly conditional", level=2)
    doc.add_paragraph(
        "Replace any wording that implies every ZIC site must publish a feed. A feed is required when the domain publishes serial content such as posts, essays, news, releases, or episodes. It is not required for a finite set of static pages."
    )
    add_bullet(doc, "When required, the RSS or Atom feed must validate and be declared with link rel=\"alternate\".")
    add_bullet(doc, "The agent-readability checklist should test the feed only when the feed requirement applies.")

    doc.add_heading("4. Expand outside-in delivery verification", level=2)
    doc.add_paragraph(
        "Part 4.4 already requires a live AI assistant to fetch and summarise a page. Version 1.2 should make the reason and the transport checks explicit: valid source is irrelevant if the public door does not open."
    )
    for item in (
        "Fetch every canonical HTTPS URL from outside the hosting account without cookies, JavaScript, or a manually accepted certificate exception.",
        "Verify that the TLS certificate is valid for the canonical hostname, not merely issued by a trusted authority.",
        "Verify that HTTP redirects to the canonical HTTPS origin only after the correct certificate is active.",
        "Fetch robots.txt, llms.txt, sitemap.xml, any applicable feed, and representative pages using at least one named crawler user-agent.",
        "Treat a host, CDN, certificate, DNS, or bot-control failure as a ZIC failure even when the repository and built HTML are correct.",
    ):
        add_bullet(doc, item)

    doc.add_heading("Proposed implementation guidance", level=1)
    doc.add_heading("Generate the agent surface from one source", level=2)
    doc.add_paragraph(
        "Strengthen Part 5 so llms.txt, llms-full.txt, Markdown mirrors, metadata, and feed entries are generated from the same content records used to build the pages. Human curation should shape descriptions and priority, but duplicated page text should not be hand-maintained."
    )

    doc.add_heading("Keep access declarations public", level=2)
    doc.add_paragraph(
        "A ZIC-Open accessibility statement should describe access for human visitors, assistive technologies, and AI agents. It should link the crawler policy and agent map, state the training position, name the licence, and explain outside-in verification."
    )

    doc.add_heading("Editorial corrections", level=1)
    for item in (
        "Use consistent British English in Charli.info-facing material, including 'colour' rather than 'color'.",
        "Keep robots.txt comments as part of the policy statement, not as disposable implementation notes.",
        "Describe the public repository as a second front door where the source is intentionally open.",
    ):
        add_bullet(doc, item)

    doc.add_heading("Acceptance checks for v1.2", level=1)
    for item in (
        "The crawler list and reference robots.txt contain the same named agents, including Meta-ExternalAgent.",
        "The focus rule protects visible keyboard focus while permitting the narrow main-landmark exception.",
        "The feed requirement has an objective applicability test.",
        "The validation section checks hostname-valid TLS and canonical HTTPS delivery from outside the host.",
        "The accessibility statement declares the Open Door and links the agent surface.",
        "Generated agent files preserve page headings and meaning without manual drift.",
    ):
        add_number(doc, item)

    doc.add_heading("Field note: why this revision exists", level=1)
    doc.add_paragraph(
        "During the Charli.info deployment, the repository, HTML, semantics, contrast, and crawler policy were sound while the canonical HTTPS origin presented a certificate for a different hostname. Conventional page-level accessibility tools could not expose that closed door because they assume the page has already loaded. The live fetch-and-summarise test did expose it. This incident validates the agent-equivalent field test in Part 4.4 and justifies making transport-level verification explicit in v1.2."
    )

    doc.add_paragraph("Prepared for Charlotte Joanne Tyrer | Draft content licensed CC BY-SA 4.0")
    path = OUT / "Zero_Information_Casualty_v1.2_Proposed_Changelog.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    for output in (build_accessibility_statement(), build_changelog()):
        print(output)
